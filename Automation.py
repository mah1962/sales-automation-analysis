# 1. Import libraries
import pandas as pd    
import matplotlib.pyplot as plt
import seaborn as sns  
import os 
from docx import Document 

# 2. Create output directory
folders=['output','charts','report'] 
for folder in folders:
    os.makedirs(folder,exist_ok=True) 
    
# 3. Load data
base_path = r"G:\python\Training Projects\P1"
sales = pd.read_csv(f"{base_path}\\sales.csv", sep=';')
customers = pd.read_csv(f"{base_path}\\customers.csv", sep=';')
costs = pd.read_csv(f"{base_path}\\costs.csv", sep=';')

# 4. Merge data
df=sales.merge(customers,on='customer_id',how='left')
df=df.merge(costs,on='product_id',how='left')
  
#  4. Simple Cleaning (NO column rename)
def smart_clean():
    for col in df.select_dtypes(include='object'):
        df[col]=df[col].str.strip()
        df[col]=df[col].dropna(how='all',inplace=True)
    return df
df.columns=df.columns.str.title()

#  5. Auto convert to numeric
def auto_convert_numeric(df,thershold=0.70):
    for col in df.select_dtypes(include='object').columns:
        converted=pd.to_numeric(df[col],errors='coerce')
        numeric_ratio=converted.notna().mean()
        if numeric_ratio>=0.7:
            df[col]=converted
    return df

#  6. Calculations
df['Revenue']=df['Quantity']*df['Unit_Price']
df['Total Cost']=df['Quantity']*df['Cost_Per_Unit']
df['Profit']=df['Revenue']-df['Total Cost']
df['Margin Profit']=(df['Profit']*100/df['Revenue']).round(0)
Max_City_Revenue=df.groupby('City')['Revenue'].sum().idxmax()
Max_Month_Revenue=df.groupby('Order_Date')['Revenue'].sum().idxmax()
Max_Category_Revenue=df.groupby('Category')['Revenue'].sum().idxmax()
Total_Revenue=df['Revenue'].sum()
Total_Orders=df['Order_Id'].nunique()
Total_Customers=df['Customer_Id'].nunique()
Average_Order_Value=Total_Revenue/Total_Orders

df
df.to_csv('G:\python\Training Projects\P1\output\sales.csv'\
    ,index=False)

#  7. KPIs
KPIs={'Total_Revenue':Total_Revenue,'Total_Orders':Total_Orders,\
    'Average_Order_Value':Average_Order_Value,\
        'Total_Customers':Total_Customers,\
        'Profit':df['Profit'].sum(),\
  'Margin Profit': df['Margin Profit'].mean() ,\
      'Max_City_Revenue':Max_City_Revenue,\
      'Max_Month_Revenue':Max_Month_Revenue,'Max_Category_Revenue'\
          :Max_Category_Revenue}

KPIs_results={'Revenue':df['Revenue'].sum(),'Total Cost':df['Total Cost'].sum(),'Profit':df['Profit'].sum(),\
  'Margin Profit': df['Margin Profit'].mean() ,'Max_City_Revenue':Max_City_Revenue,\
      'Max_Month_Revenue':Max_Month_Revenue,'Max_Category_Revenue'\
          :Max_Category_Revenue}

KPIs_df=pd.DataFrame(list(KPIs_results.items()),columns=['KPIs','Value'])
KPIs_df.to_csv(r'G:\python\Training Projects\P1\output\KPIs.csv'\
    , index=False)


# 8. Charts
mydata=df.groupby(['Quantity','Product_Name','Category'],\
    as_index=False)\
    [['Revenue','Total Cost','Profit','Margin Profit']].sum()
plt.figure(figsize=(10,6))   
Quantity_Revenue_Sum=sns.histplot(x='Quantity',y='Revenue',data=mydata)
plt.savefig('charts/Quantity_Revenue_Sum.png')
plt.close()

plt.figure(figsize=(10,6))
Quantity_Profit_Sum=sns.histplot(x='Quantity',y='Profit',data=mydata)
plt.savefig('charts/Quantity_Profit_Sum.png')
plt.close()

plt.figure(figsize=(10,6))
Product_Name_Margin=sns.histplot(x='Product_Name',y='Margin Profit'\
    ,data=mydata)
plt.savefig('charts/Product_Name_Margin.png')
plt.close()

plt.figure(figsize=(10,6))
Category_Profit_Sum=sns.histplot(x='Category',y='Profit',data=mydata)
plt.savefig('charts/Category_Profit_Sum.png')
plt.close()

plt.figure(figsize=(10,6))
Product_Name_Revenue_Sum=sns.histplot(x='Product_Name',y='Revenue'\
    ,data=mydata)
plt.savefig('charts/Product_Name_Revenue_Sum.png')
plt.close()

plt.figure(figsize=(10,6))
Category_Revenue_Sum=sns.histplot(x='Category',y='Revenue',data=mydata)
plt.savefig('charts/Category_Revenue_Sum.png')
plt.close()


# 9. INSIGHTS
Top_Category=df.groupby('Category')['Revenue'].sum().idxmax()
Top_Month=df.groupby('Order_Date')['Revenue'].sum().idxmax()

Insights = [
    f"Highest revenue category is {Top_Category}.",
    f"Peak sales month is {Top_Month}.",
    "Revenue concentration indicates strong performance in specific product categories.",
    "Average order value suggests stable customer purchasing behavior.",
]

# 10. REPORT (WORD)
doc=Document()
doc.add_heading("Sales Analysis Report",level=1)

doc.add_heading("Project Oveview",level=2)
doc.add_paragraph("This automated report analizes sales performance,\
customer behavior, and revenue trends using python automation ")

doc.add_heading("Key Performance Indicators(KPIs)",level=2)
doc.add_paragraph(f"Total Revenue : {Total_Revenue:,.2f}")
doc.add_paragraph(f"Total Orders : {Total_Orders}")
doc.add_paragraph(f"Average Order Value : {Average_Order_Value}")
doc.add_paragraph(f"Total_Customers : {Total_Customers:,.2f}")

doc.add_heading("Visual Analysis",level=2)
doc.add_picture("charts/Category_Profit_Sum.png")
doc.add_picture("charts/Category_Revenue_Sum.png")
doc.add_picture("charts/Product_Name_Margin.png")
doc.add_picture("charts/Product_Name_Revenue_Sum.png")
doc.add_picture("charts/Quantity_Profit_Sum.png")
doc.add_picture("charts/Quantity_Revenue_Sum.png")

doc.add_heading("Conclusion",level=2)
doc.add_paragraph(
    "The analysis highlights key revenue drivers and sales patterns. "
    "These insights can support data-driven business decisions." ) 

doc.save("report/Sales_Analysis_Report.docx")

# 11. Done
import logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s\
    - %(message)s')

print("Automation completed successfuly!")
print("Outputs saved in: output/charts/report")







 


    




